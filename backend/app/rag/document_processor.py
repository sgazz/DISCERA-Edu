"""
Advanced Document Processing Pipeline for DISCERA RAG System
"""
import os
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
import PyPDF2
from docx import Document
import re
from dataclasses import dataclass

from app.core.config import settings

logger = logging.getLogger(__name__)


@dataclass
class DocumentChunk:
    """Document chunk with metadata"""
    content: str
    chunk_id: str
    page_number: Optional[int] = None
    section: Optional[str] = None
    metadata: Dict[str, Any] = None


class DocumentProcessor:
    """Advanced document processor for multiple formats"""
    
    def __init__(self):
        self.supported_extensions = {'.pdf', '.docx', '.doc', '.txt'}
        self.chunk_size = 1000  # characters
        self.chunk_overlap = 200  # characters
    
    def process_document(self, file_path: str, file_type: str) -> List[DocumentChunk]:
        """Process document and return chunks"""
        try:
            if file_type.lower() == '.pdf':
                return self._process_pdf(file_path)
            elif file_type.lower() in ['.docx', '.doc']:
                return self._process_docx(file_path)
            elif file_type.lower() == '.txt':
                return self._process_txt(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}")
            raise
    
    def _process_pdf(self, file_path: str) -> List[DocumentChunk]:
        """Process PDF document"""
        chunks = []
        page_number = 1
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text.strip():
                        # Clean text
                        text = self._clean_text(text)
                        
                        # Split into chunks
                        page_chunks = self._create_chunks(text, page_number)
                        chunks.extend(page_chunks)
                    
                    page_number += 1
                
                logger.info(f"Processed PDF: {len(chunks)} chunks from {page_number-1} pages")
                return chunks
                
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            raise
    
    def _process_docx(self, file_path: str) -> List[DocumentChunk]:
        """Process DOCX document"""
        chunks = []
        
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            # Combine all text
            full_text = "\n".join(text_parts)
            full_text = self._clean_text(full_text)
            
            # Create chunks
            chunks = self._create_chunks(full_text)
            
            logger.info(f"Processed DOCX: {len(chunks)} chunks")
            return chunks
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            raise
    
    def _process_txt(self, file_path: str) -> List[DocumentChunk]:
        """Process TXT document"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            text = self._clean_text(text)
            chunks = self._create_chunks(text)
            
            logger.info(f"Processed TXT: {len(chunks)} chunks")
            return chunks
            
        except Exception as e:
            logger.error(f"Error processing TXT {file_path}: {e}")
            raise
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep important ones
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)\[\]\{\}]', '', text)
        
        # Normalize quotes
        text = text.replace('"', '"').replace('"', '"')
        text = text.replace(''', "'").replace(''', "'")
        
        return text.strip()
    
    def _create_chunks(self, text: str, page_number: Optional[int] = None) -> List[DocumentChunk]:
        """Create semantic chunks from text"""
        chunks = []
        
        # Split by sentences first
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if s.strip()]
        
        current_chunk = ""
        chunk_id = 0
        
        for sentence in sentences:
            # If adding this sentence would exceed chunk size
            if len(current_chunk) + len(sentence) > self.chunk_size and current_chunk:
                # Save current chunk
                chunk = DocumentChunk(
                    content=current_chunk.strip(),
                    chunk_id=f"chunk_{chunk_id}",
                    page_number=page_number,
                    metadata={
                        "chunk_size": len(current_chunk),
                        "sentence_count": current_chunk.count('.') + 1
                    }
                )
                chunks.append(chunk)
                
                # Start new chunk with overlap
                overlap_text = current_chunk[-self.chunk_overlap:] if self.chunk_overlap > 0 else ""
                current_chunk = overlap_text + " " + sentence
                chunk_id += 1
            else:
                current_chunk += " " + sentence
        
        # Add final chunk
        if current_chunk.strip():
            chunk = DocumentChunk(
                content=current_chunk.strip(),
                chunk_id=f"chunk_{chunk_id}",
                page_number=page_number,
                metadata={
                    "chunk_size": len(current_chunk),
                    "sentence_count": current_chunk.count('.') + 1
                }
            )
            chunks.append(chunk)
        
        return chunks
    
    def get_document_summary(self, chunks: List[DocumentChunk]) -> Dict[str, Any]:
        """Generate document summary from chunks"""
        total_chunks = len(chunks)
        total_content_length = sum(len(chunk.content) for chunk in chunks)
        
        # Extract key information
        pages = set(chunk.page_number for chunk in chunks if chunk.page_number)
        
        return {
            "total_chunks": total_chunks,
            "total_content_length": total_content_length,
            "pages": len(pages),
            "average_chunk_size": total_content_length / total_chunks if total_chunks > 0 else 0
        } 